# -*- coding: utf-8 -*-
"""
Created on Tue Jul 30 15:59:17 2019

@author: 无敌钢牙小白狼
"""

import pandas as pd
import os
import sqlite3  

open_path = 'open_path/'
save_path = '本地数据库/'
os.chdir('C:\\Users\\无敌钢牙小白狼\\Desktop\\德勤数旅\\数据处理')
#%% 读取文件对应建立本地数据库
''' 数据库名字 日期+data  '''
df = pd.read_excel(open_path+'tagdata.xls')


conn = sqlite3.connect(save_path+'data.db')
cur = conn.cursor()
df.to_sql('20190730data',conn,if_exists='replace')
conn.close
#%%基金经理数据处理存到数据库
''' 数据库名字 日期+manager table 是季度数据 '''
conn = sqlite3.connect(save_path+'manager.db')
cur = conn.cursor()

manager = pd.ExcelFile(open_path+'fundmanager.xlsx')
sheetnames = manager.sheet_names

for i in sheetnames:
    print(i)
    df_temp = pd.read_excel(open_path+'fundmanager.xlsx',sheet_name = i)
    name = i
    df_temp.to_sql(i,conn,if_exists='replace')
conn.close
#%%读取本地数据库
conn = sqlite3.connect(save_path+'data.db')
cur = conn.cursor()

sql = ''' select * from '20190730data' '''
df = pd.read_sql(sql,con=conn)

conn.close

conn = sqlite3.connect(save_path+'manager.db')
sql = ''' select * from '20170930' '''
df_manager1 = pd.read_sql(sql,con = conn)
sql = ''' select * from '20170630' '''
df_manager2 = pd.read_sql(sql,con = conn)
sql = ''' select * from '20170331' '''
df_manager3 = pd.read_sql(sql,con = conn)

conn.close
#%%处理数据
def handle_first(df_manager):
    dic = {}
    for i in df_manager.index:
        df_temp = df_manager.loc[i,:]    
        temp = df_temp.sort_values()[:3]
        temp = list(temp.index)
        perfer_list = []
        for industry in temp:
            perfer_list.append(industry.split('(',1)[0])
        dic[i] = perfer_list    
    return dic

def handle_next(df_manager,dic):
    for i in df_manager.index:
        df_temp = df_manager.loc[i,:]    
        temp = df_temp.sort_values()[:3]
        temp = list(temp.index)
        perfer_list = []
        for industry in temp:
            perfer_list.append(industry.split('(',1)[0])    
        if i in dic:  
            perfer_list_f = perfer_list + dic[i]
            perfer_list_f = list(set(perfer_list_f))
            dic[i] = perfer_list_f
        else:
            dic[i] = perfer_list   
    return dic



dic = handle_first(df_manager1)
dic = handle_next(df_manager2,dic)
dic = handle_next(df_manager3,dic)

#%%匹配信息
manager = input('输入基金经理：   ')
industry = dic[manager]

dic_target = {}
for i in industry:
    df_temp = df[df.所属申万一级行业 == i]
    dic_target[i] = df_temp
    
#%%每个基金经理生成一个
#for i in dic:
#    if not os.path.isdir(report_path+i):
#       os.makedirs(report_path+i)



#%%生成文件夹
report_path = '报告输出/'
os.makedirs(report_path+manager)  
for i in dic_target:
    temp = dic_target[i]
    if len(temp)>0:
        os.makedirs(report_path+manager+'/'+i)        
#%%生成报告
from docx import Document


for i in dic_target:
    temp = dic_target[i]
    if len(temp)>0:
        report_path_f =  report_path+manager+'/'+i+'/'
        for name in temp['名称']:
            Doc = Document() 
            df_temp = temp[temp.名称==name]
            Name = Doc.add_heading(name+'报告速递',level=1)

            
            stock_info = Doc.add_heading('\n个股信息:',level=2)
            sql = '''  
               \t股票上市代码为： %s \t 股票申购代码 %s  \t 上市时间为：  %s \t股票发行价格：%s \t \
            共有 %s 家机构做出股价预测, 预测值为 %s  元/股         
            '''%(df_temp.代码.values, df_temp.申购代码.values, df_temp.上市日期.values, df_temp.发行价格.values,\
            df_temp.研报预测个数.values, df_temp.研报预测价格.values)
            stock_info.add_run(sql)

            
            stock_news = Doc.add_heading('\n个股新闻',level=2)
            sql = '''
            \t 收集个股新闻20个，与股东相关3个，行业相关1个，有效信息 %s 个，其中正面新闻 %s 个，负面新闻 %s 个
            '''%(df_temp.正面新闻.values+df_temp.负面新闻.values,df_temp.正面新闻.values,df_temp.负面新闻.values)
            stock_news.add_run(sql)

            
            industry = Doc.add_heading('\n所属行业分析',level=2)
            sql = '''
            \t 所属申万一级行业 %s , 行业近一个月PE %s, 行业有利信息： %s:
            '''%(df_temp.所属申万一级行业.values,df_temp['行业PE(近1月,TTM)'].values,df_temp.所属行业有利信息.values )
            industry.add_run(sql)

      
            
            Doc.save(report_path_f+name+'.docx')
            











